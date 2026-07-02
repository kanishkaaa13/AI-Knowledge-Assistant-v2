from __future__ import annotations

import hashlib
import re
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from cryptography.fernet import InvalidToken
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.core.crypto import encryption_service
from app.core.sanitize import ensure_present, sanitize_text
from app.models.uploaded_document import UploadedDocument
from app.models.user import User
from app.repositories.document import DocumentRepository


MAX_PREVIEW_LENGTH = 2400


@dataclass
class ExtractedDocumentData:
    file_extension: str
    mime_type: str | None
    file_size: int
    checksum: str
    extracted_text: str
    page_count: int | None
    word_count: int
    storage_path: str
    safe_file_name: str
    encrypted_size: int


def _sanitize_file_name(file_name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", file_name).strip("._-")
    return cleaned or "document"


def _extract_text(file_extension: str, file_bytes: bytes) -> tuple[str, int | None]:
    if file_extension in {".txt", ".md"}:
        try:
            return file_bytes.decode("utf-8"), None
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1"), None

    if file_extension == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, len(reader.pages)

    if file_extension == ".docx":
        document = DocxDocument(BytesIO(file_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return text, len(document.paragraphs)

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST,
        detail="Unsupported file type.",
    )


def _resolve_upload_dir(user_id: uuid.UUID) -> Path:
    upload_root = Path(settings.UPLOAD_ROOT_DIR)
    user_dir = upload_root / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir


async def parse_upload(file: UploadFile, user_id: uuid.UUID) -> ExtractedDocumentData:
    original_name = file.filename or "document"
    safe_name = _sanitize_file_name(original_name)
    extension = Path(safe_name).suffix.lower()

    if extension not in settings.ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Unsupported file type. Allowed: PDF, DOCX, TXT, MD.",
        )

    file_bytes = await file.read()
    file_size = len(file_bytes)
    if file_size == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty.",
        )

    if file_size > settings.MAX_UPLOAD_SIZE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File exceeds maximum size of {settings.MAX_UPLOAD_SIZE_BYTES // (1024 * 1024)} MB.",
        )

    checksum = hashlib.sha256(file_bytes).hexdigest()
    extracted_text, page_count = _extract_text(extension, file_bytes)
    word_count = len(extracted_text.split()) if extracted_text else 0

    allowed_mime_types = {
        ".pdf": {"application/pdf"},
        ".docx": {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/zip",
        },
        ".txt": {"text/plain"},
        ".md": {"text/markdown", "text/plain", "text/x-markdown"},
    }
    if file.content_type and file.content_type not in allowed_mime_types.get(extension, set()):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File content type does not match the selected extension.",
        )

    user_dir = _resolve_upload_dir(user_id)
    stored_name = f"{checksum[:16]}-{safe_name}.bin"
    file_path = user_dir / stored_name
    encrypted_bytes = encryption_service.encrypt_bytes(file_bytes)
    file_path.write_bytes(encrypted_bytes)

    return ExtractedDocumentData(
        file_extension=extension,
        mime_type=file.content_type,
        file_size=file_size,
        checksum=checksum,
        extracted_text=extracted_text,
        page_count=page_count,
        word_count=word_count,
        storage_path=str(file_path.resolve()),
        safe_file_name=safe_name,
        encrypted_size=len(encrypted_bytes),
    )


def preview_text(text: str | None) -> str | None:
    if not text:
        return None
    return text[:MAX_PREVIEW_LENGTH]


def create_document_record(
    db: Session,
    *,
    user: User,
    title: str,
    upload_data: ExtractedDocumentData,
) -> UploadedDocument:
    repository = DocumentRepository(db)
    safe_title = ensure_present(sanitize_text(title, max_length=255), field_name="title")
    existing = repository.get_by_user_and_checksum(user.id, upload_data.checksum)
    if existing:
        existing_path = Path(upload_data.storage_path)
        if existing_path.exists():
            existing_path.unlink()
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="This document has already been uploaded.",
        )

    return repository.create(
        user_id=user.id,
        title=safe_title,
        file_name=upload_data.safe_file_name,
        file_extension=upload_data.file_extension,
        file_path=upload_data.storage_path,
        mime_type=upload_data.mime_type,
        file_size=upload_data.file_size,
        checksum=upload_data.checksum,
        page_count=upload_data.page_count,
        word_count=upload_data.word_count,
        status="processed",
        extracted_text=upload_data.extracted_text,
    )


def delete_document_file(document: UploadedDocument) -> None:
    if not document.file_path:
        return
    file_path = Path(document.file_path)
    if file_path.exists():
        file_path.unlink()


def read_encrypted_document_bytes(document: UploadedDocument) -> bytes:
    if not document.file_path:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found.")

    file_path = Path(document.file_path)
    if not file_path.exists():
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Stored file not found.")

    try:
        return encryption_service.decrypt_bytes(file_path.read_bytes())
    except InvalidToken as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Stored file could not be decrypted.",
        ) from exc
