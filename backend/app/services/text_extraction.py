"""Shared plain-text extraction helpers for uploaded documents."""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfReader


class UnsupportedFileTypeError(ValueError):
    """Raised when a file extension has no text extractor."""

    def __init__(self, file_extension: str) -> None:
        super().__init__(f"Unsupported file type: {file_extension}")
        self.file_extension = file_extension


def extract_text(file_extension: str, file_bytes: bytes) -> tuple[str, int | None]:
    """Extract text from raw file bytes.

    Returns a tuple of (extracted_text, page_count) where page_count is None for
    formats without pagination.

    Raises:
        UnsupportedFileTypeError: If the extension has no extractor.
    """
    if file_extension in {".txt", ".md"}:
        try:
            return file_bytes.decode("utf-8"), None
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1"), None

    if file_extension == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text, len(reader.pages)

    if file_extension == ".docx":
        document = DocxDocument(BytesIO(file_bytes))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return text, len(document.paragraphs)

    raise UnsupportedFileTypeError(file_extension)
