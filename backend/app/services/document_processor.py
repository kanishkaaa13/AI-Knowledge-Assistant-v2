from __future__ import annotations

import hashlib
import re
import uuid
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.document_chunk import DocumentChunk
from app.models.uploaded_document import UploadedDocument


@dataclass
class ProcessedDocument:
    document_id: uuid.UUID
    chunk_count: int
    file_path: str
    file_size: int
    file_type: str


class DocumentProcessor:
    """Service for processing uploaded documents: text extraction, chunking, and database storage."""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text(self, file_bytes: bytes, file_extension: str) -> tuple[str, int | None]:
        """
        Extract text from file bytes based on file type.

        Args:
            file_bytes: Raw file content as bytes
            file_extension: File extension (e.g., .pdf, .docx, .txt, .md)

        Returns:
            Tuple of (extracted_text, page_count or None)

        Raises:
            ValueError: If file type is unsupported
        """
        if file_extension in {".txt", ".md"}:
            try:
                return file_bytes.decode("utf-8"), None
            except UnicodeDecodeError:
                return file_bytes.decode("latin-1"), None

        if file_extension == ".pdf":
            reader = PdfReader(BytesIO(file_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text, len(reader.pages)

        if file_extension == ".docx":
            document = DocxDocument(BytesIO(file_bytes))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            return text, len(document.paragraphs)

        raise ValueError(f"Unsupported file type: {file_extension}")

    def chunk_text(self, text: str) -> list[str]:
        """
        Split text into chunks with intelligent sentence boundary detection.

        Args:
            text: The text to chunk

        Returns:
            List of text chunks
        """
        if not text or not text.strip():
            return []

        # Normalize whitespace
        text = re.sub(r"\s+", " ", text.strip())

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            # Calculate end position
            end = start + self.chunk_size

            # If we're not at the end, try to find a good break point
            if end < text_length:
                # Look for sentence boundaries (., !, ?) followed by space
                for i in range(end, max(start + self.chunk_size // 2, start), -1):
                    if text[i] in ".!?" and i + 1 < text_length and text[i + 1] == " ":
                        end = i + 1
                        break
                else:
                    # If no sentence boundary, look for word boundary
                    for i in range(end, max(start + self.chunk_size // 2, start), -1):
                        if text[i] == " ":
                            end = i
                            break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            # Move start position with overlap
            start = end - self.chunk_overlap
            if start < 0:
                start = 0

        return chunks

    def save_file(self, file_bytes: bytes, user_id: uuid.UUID, file_name: str) -> tuple[str, str]:
        """
        Save file to user-isolated storage directory.

        Args:
            file_bytes: Raw file content
            user_id: User ID for isolation
            file_name: Original file name

        Returns:
            Tuple of (file_path, checksum)
        """
        upload_root = Path(settings.UPLOAD_ROOT_DIR)
        user_dir = upload_root / str(user_id)
        user_dir.mkdir(parents=True, exist_ok=True)

        checksum = hashlib.sha256(file_bytes).hexdigest()
        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "-", file_name).strip("._-")
        stored_name = f"{checksum[:16]}-{safe_name}"
        file_path = user_dir / stored_name

        file_path.write_bytes(file_bytes)
        return str(file_path.resolve()), checksum

    def create_document_record(
        self,
        db: Session,
        user_id: uuid.UUID,
        title: str,
        file_name: str,
        file_extension: str,
        file_path: str,
        file_size: int,
        checksum: str,
        extracted_text: str,
        page_count: int | None = None,
        mime_type: str | None = None,
    ) -> UploadedDocument:
        """
        Create database record for uploaded document.

        Args:
            db: Database session
            user_id: User ID
            title: Document title
            file_name: Original file name
            file_extension: File extension
            file_path: Storage file path
            file_size: File size in bytes
            checksum: SHA256 checksum
            extracted_text: Extracted text content
            page_count: Page count (for PDFs)
            mime_type: MIME type

        Returns:
            Created UploadedDocument instance
        """
        word_count = len(extracted_text.split()) if extracted_text else 0

        document = UploadedDocument(
            user_id=user_id,
            title=title,
            file_name=file_name,
            file_extension=file_extension,
            file_path=file_path,
            mime_type=mime_type,
            file_size=file_size,
            checksum=checksum,
            page_count=page_count,
            word_count=word_count,
            status="processed",
            extracted_text=extracted_text,
        )

        db.add(document)
        db.commit()
        db.refresh(document)

        return document

    def create_chunk_records(
        self,
        db: Session,
        document: UploadedDocument,
        chunks: list[str],
    ) -> list[DocumentChunk]:
        """
        Create database records for document chunks.

        Args:
            db: Database session
            document: Parent document
            chunks: List of text chunks

        Returns:
            List of created DocumentChunk instances
        """
        chunk_records = []

        for index, chunk_content in enumerate(chunks):
            chunk = DocumentChunk(
                document_id=document.id,
                chunk_index=index,
                page_number=document.page_count if document.page_count else None,
                content=chunk_content,
                token_count=len(chunk_content.split()),
            )
            chunk_records.append(chunk)
            db.add(chunk)

        db.commit()

        for chunk in chunk_records:
            db.refresh(chunk)

        return chunk_records

    def process_document(
        self,
        db: Session,
        user_id: uuid.UUID,
        title: str,
        file_name: str,
        file_bytes: bytes,
        mime_type: str | None = None,
    ) -> ProcessedDocument:
        """
        Complete document processing pipeline: save, extract, chunk, and store.

        Args:
            db: Database session
            user_id: User ID
            title: Document title
            file_name: Original file name
            file_bytes: Raw file content
            mime_type: MIME type (optional)

        Returns:
            ProcessedDocument with document_id, chunk_count, and metadata

        Raises:
            ValueError: If file type is unsupported
        """
        file_extension = Path(file_name).suffix.lower()

        # Validate file type
        if file_extension not in settings.ALLOWED_UPLOAD_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_extension}. Allowed: {settings.ALLOWED_UPLOAD_EXTENSIONS}"
            )

        # Validate file size
        file_size = len(file_bytes)
        if file_size > settings.MAX_UPLOAD_SIZE_BYTES:
            raise ValueError(
                f"File exceeds maximum size of {settings.MAX_UPLOAD_SIZE_BYTES // (1024 * 1024)} MB"
            )

        # Save file to storage
        file_path, checksum = self.save_file(file_bytes, user_id, file_name)

        # Extract text
        extracted_text, page_count = self.extract_text(file_bytes, file_extension)

        # Create document record
        document = self.create_document_record(
            db=db,
            user_id=user_id,
            title=title,
            file_name=file_name,
            file_extension=file_extension,
            file_path=file_path,
            file_size=file_size,
            checksum=checksum,
            extracted_text=extracted_text,
            page_count=page_count,
            mime_type=mime_type,
        )

        # Chunk text
        chunks = self.chunk_text(extracted_text)

        # Create chunk records
        chunk_records = self.create_chunk_records(db, document, chunks)

        return ProcessedDocument(
            document_id=document.id,
            chunk_count=len(chunk_records),
            file_path=file_path,
            file_size=file_size,
            file_type=file_extension,
        )
